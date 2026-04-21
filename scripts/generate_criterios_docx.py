"""Generate a clean .docx presentation document for the Criterios template
that can be sent directly to writers."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = r"C:\Users\Usuario\eplus-tools\docs\CRITERIOS_TEMPLATE.docx"

NAVY = RGBColor(0x1B, 0x14, 0x64)
YELLOW = RGBColor(0xE7, 0xEB, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY_BG = "F4F4F4"
DARK = RGBColor(0x22, 0x22, 0x22)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 22, 2: 16, 3: 13}
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=DARK, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    run = p.runs[0] if p.runs else p.add_run("")
    p.text = ""
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), LIGHT_GREY_BG)
    pPr.append(shd)
    return p


def add_callout(doc, label, text, color=YELLOW):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8DC" if color == YELLOW else "F0F4FF")
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    label_run = p.add_run(label + " ")
    label_run.bold = True
    label_run.font.color.rgb = NAVY
    label_run.font.size = Pt(11)
    text_run = p.add_run(text)
    text_run.font.size = Pt(11)
    text_run.font.color.rgb = DARK
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], "1B1464")
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
    return table


# ─── Build document ─────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Cover ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("PLANTILLA DE CRITERIOS DE EVALUACIÓN")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY
cover.paragraph_format.space_before = Pt(120)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Guía para escribas — qué entregar y cómo")
run.font.size = Pt(14)
run.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(40)

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = intro.add_run(
    "Por cada pregunta del formulario, entregas un documento\n"
    "con un bloque genérico + entre 4 y 6 criterios."
)
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = DARK

doc.add_page_break()

# ── Section 1: Resumen ──
add_heading(doc, "1. Resumen de una línea", level=1)
add_callout(
    doc,
    "EN UNA FRASE:",
    "Por cada pregunta entregas un bloque genérico (5 campos) + "
    "entre 4 y 6 criterios, cada uno con 5 campos, todo en prosa natural.",
)

# ── Section 2: Estructura ──
add_heading(doc, "2. Estructura del documento por pregunta", level=1)
add_para(doc, "Cada documento tiene dos partes:", size=11)

add_code_block(
    doc,
    "═══════════════════════════════════════════════════════════\n"
    "PARTE A — BLOQUE GENÉRICO DE LA PREGUNTA\n"
    "═══════════════════════════════════════════════════════════\n"
    "  1. IDENTIFICACIÓN\n"
    "  2. CONTEXTO DE LA PREGUNTA\n"
    "  3. CONEXIONES (apoya en / alimenta a)\n"
    "  4. REGLA GLOBAL (opcional)\n"
    "  5. LÍMITE DE EXTENSIÓN\n"
    "\n"
    "═══════════════════════════════════════════════════════════\n"
    "PARTE B — CRITERIOS (entre 4 y 6)\n"
    "═══════════════════════════════════════════════════════════\n"
    "  CRITERIO 1\n"
    "    Título corto + Prioridad + Obligatorio\n"
    "    INTENCIÓN\n"
    "    ELEMENTOS\n"
    "    EJEMPLOS (débil + fuerte)\n"
    "    EVITAR\n"
    "\n"
    "  CRITERIO 2 ...\n"
    "  CRITERIO 3 ..."
)

doc.add_page_break()

# ── Section 3: Parte A ──
add_heading(doc, "3. PARTE A — Bloque genérico de la pregunta", level=1)
add_para(doc, "Estos 5 campos van una sola vez al inicio del documento.", italic=True, color=GREY)

add_heading(doc, "1. IDENTIFICACIÓN", level=2)
add_para(doc, "Código y título oficial de la pregunta tal y como aparece en el formulario.")
add_callout(doc, "EJEMPLO:", "1.1 Background and general objectives")

add_heading(doc, "2. CONTEXTO DE LA PREGUNTA", level=2)
add_para(
    doc,
    "Un párrafo (4-6 frases) que explica de qué va esta pregunta en esencia y "
    "qué busca el evaluador. Es el marco mental antes de leer los criterios."
)
add_callout(
    doc,
    "EJEMPLO:",
    'Esta pregunta pide dos cosas a la vez: explicar el problema o necesidad que el '
    'proyecto aborda, y demostrar que el proyecto encaja con el propósito y prioridades '
    'de la convocatoria. Es una de las más importantes — incluso una buena idea recibe '
    'puntuación baja si la relevancia respecto a la call no queda clara. El evaluador '
    'chequea: ¿es un problema real? ¿a quién afecta? ¿por qué importa? ¿qué cambio creará '
    'el proyecto? ¿encaja con esta call concreta?'
)

add_heading(doc, "3. CONEXIONES", level=2)
add_para(doc, "Dos sub-bloques narrativos (no bullets pelados):")
add_bullet(doc, "APOYA EN: qué preguntas previas son fundamento de ésta. Si es la primera, escribir 'Es la pregunta inicial — no hay secciones previas'.")
add_bullet(doc, "ALIMENTA A: qué preguntas posteriores se construyen sobre ésta.")
add_callout(
    doc,
    "EJEMPLO:",
    "APOYA EN: Es la pregunta inicial. ALIMENTA A: 1.2 profundizará el problema con datos "
    "cuantitativos. 2.1.1 (metodología) debe responder a estos objetivos. 3.1 (impacto) "
    "debe medir el cambio prometido."
)

add_heading(doc, "4. REGLA GLOBAL (opcional)", level=2)
add_para(
    doc,
    "Una o dos frases que enuncian un principio transversal aplicable a todos los "
    "criterios de esta pregunta. NO siempre existe — solo en casos como Small-scale "
    "Partnerships."
)
add_callout(
    doc,
    "EJEMPLO:",
    "Small-scale logic: scope focalizado, objetivos limitados, ambición proporcional al "
    "budget y duración."
)

add_heading(doc, "5. LÍMITE DE EXTENSIÓN", level=2)
add_para(doc, "Palabras o páginas máximas para la respuesta final.")
add_callout(doc, "EJEMPLO:", "Max 3 páginas / ~1500 palabras")

doc.add_page_break()

# ── Section 4: Parte B ──
add_heading(doc, "4. PARTE B — Criterios", level=1)
add_para(
    doc,
    "Entre 4 y 6 criterios por pregunta. Cada criterio = un párrafo / idea grande "
    "que el evaluador chequea en la respuesta."
)
add_callout(
    doc,
    "REGLA DE ORO:",
    "Si te salen 7+ criterios, fusiona. Si te salen menos de 4, probablemente la pregunta "
    "no es tan compleja como crees."
)

add_para(doc, "Cada criterio tiene 5 campos:", bold=True, size=12)

add_heading(doc, "Metadatos del criterio", level=2)
add_bullet(doc, "Título corto (5-7 palabras): etiqueta del criterio.")
add_bullet(doc, "Prioridad: alta / media / baja.")
add_bullet(doc, "Obligatorio: sí / no.")

add_heading(doc, "INTENCIÓN", level=2)
add_para(doc, "2-3 frases que explican qué debe demostrar este párrafo. Es el 'para qué' del criterio.")

add_heading(doc, "ELEMENTOS", level=2)
add_para(doc, "Qué tiene que aparecer concretamente en el texto. Puede ser una breve lista o prosa.")

add_heading(doc, "EJEMPLOS — el más importante", level=2)
add_bullet(doc, "DÉBIL: 1-2 frases reales que un mal escritor pondría.")
add_bullet(doc, "FUERTE: 1-2 frases que un buen escritor pondría.")
add_callout(
    doc,
    "POR QUÉ IMPORTA TANTO:",
    "La IA aprende mejor con un contraste concreto que con una descripción abstracta. "
    "Sin estos ejemplos, los criterios pierden el 80% de su valor."
)

add_heading(doc, "EVITAR", level=2)
add_para(doc, "3-4 errores típicos que penalizan. Específicos para este criterio.")

doc.add_page_break()

# ── Section 5: Reglas de oro ──
add_heading(doc, "5. Reglas de oro al escribir", level=1)
reglas = [
    "Prosa natural, no checklist. Escribe como si le explicaras por chat a alguien junior. Evita bullets pelados sin contexto.",
    "Los ejemplos débil/fuerte son obligatorios. Aunque te cueste, inventa uno realista. Sin ejemplos no hay brief útil.",
    "Sé específica. 'Frases genéricas sobre el deporte' es vago. Mejor: 'Empezar con: El deporte juega un papel fundamental en la sociedad'.",
    "Máximo 6 criterios por pregunta. Si salen más, fusiona los que se solapen.",
    "Cada criterio = un párrafo / idea grande del evaluador. Si dudas si dos cosas son criterios distintos, pregúntate: ¿el evaluador las puntúa por separado o juntas? Si juntas → un criterio.",
    "No copies del Programme Guide. Reformula con tus palabras y conecta a la realidad del proyecto.",
    "Si una idea es transversal (aplica a TODOS los criterios), va en REGLA GLOBAL no como criterio aparte.",
]
for i, r in enumerate(reglas, 1):
    p = doc.add_paragraph()
    num_run = p.add_run(f"{i}. ")
    num_run.bold = True
    num_run.font.color.rgb = NAVY
    num_run.font.size = Pt(11)
    text_run = p.add_run(r)
    text_run.font.size = Pt(11)
    text_run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ── Section 6: Ejemplo completo ──
add_heading(doc, "6. Ejemplo completo — Pregunta 1.1 bien hecha", level=1)
add_para(
    doc,
    "Esto es un documento de criterios totalmente desarrollado. Es el modelo a seguir.",
    italic=True, color=GREY,
)

add_heading(doc, "PARTE A — Bloque genérico", level=2)

add_para(doc, "1. IDENTIFICACIÓN", bold=True, color=NAVY)
add_para(doc, "1.1 Background and general objectives")

add_para(doc, "2. CONTEXTO DE LA PREGUNTA", bold=True, color=NAVY)
add_para(
    doc,
    "Esta pregunta pide dos cosas a la vez: explicar el problema o necesidad que el "
    "proyecto aborda, y demostrar que el proyecto encaja con el propósito y prioridades "
    "de la convocatoria. Es una de las preguntas más importantes — incluso una buena "
    "idea recibe puntuación baja si la relevancia respecto a la call no queda clara. El "
    "evaluador chequea cinco cosas: ¿es un problema real?, ¿a quién afecta?, ¿por qué "
    "importa?, ¿qué cambio creará el proyecto? y ¿encaja con esta call concreta?"
)

add_para(doc, "3. CONEXIONES", bold=True, color=NAVY)
add_para(doc, "APOYA EN: Es la pregunta inicial — no hay secciones previas que la sustenten.")
add_para(
    doc,
    "ALIMENTA A: 1.2 (Needs analysis) profundizará el problema con datos cuantitativos e "
    "indicadores. 2.1.1 (Methodology) debe responder directamente a los objetivos generales "
    "definidos aquí. 3.1 (Impact) debe medir el cambio prometido en estos objetivos."
)

add_para(doc, "4. REGLA GLOBAL", bold=True, color=NAVY)
add_para(
    doc,
    "Small-scale logic: scope focalizado, objetivos limitados, ambición proporcional al "
    "budget y duración."
)

add_para(doc, "5. LÍMITE DE EXTENSIÓN", bold=True, color=NAVY)
add_para(doc, "Max 3 páginas / ~1500 palabras.")

add_heading(doc, "PARTE B — Criterios (5)", level=2)

criterios_ejemplo = [
    {
        "titulo": "Background grounded in a real problem",
        "prioridad": "alta",
        "obligatorio": "sí",
        "intencion": (
            "El párrafo debe abrir describiendo la situación que llevó a la idea del "
            "proyecto. El foco está en el PROBLEMA, no en lo que vais a hacer. El "
            "evaluador tiene que entender en 30 segundos cuál es el problema y por qué "
            "merece atención."
        ),
        "elementos": [
            "Sector y entorno geográfico donde operan las organizaciones (local, regional, grassroots).",
            "Un problema principal claramente formulado (no varios problemas mezclados).",
            "Quién está directamente afectado (target groups concretos, no 'la sociedad').",
            "Por qué importa: consecuencias si no se actúa.",
        ],
        "ejemplo_debil": (
            "Sport plays a fundamental role in European society, and there is a need to "
            "strengthen its impact on young people."
        ),
        "ejemplo_fuerte": (
            "In small grassroots sport clubs across northern Spain, over 70% of coaches are "
            "volunteers without formal pedagogical training. This contributes to early "
            "dropout among adolescents aged 12-16, particularly girls, who leave competitive "
            "sport after negative experiences in their first years of participation."
        ),
        "evitar": [
            "Empezar con 'This project aims to...'",
            "Frases genéricas sobre la importancia del deporte",
            "Largas explicaciones sobre Erasmus+ o la convocatoria",
            "Mezclar varios problemas no relacionados",
        ],
    },
    {
        "titulo": "Justificación con evidencia",
        "prioridad": "alta",
        "obligatorio": "sí",
        "intencion": (
            "Tras presentar el problema, demostrar que existe un GAP real (algo está "
            "faltando o no funciona) y respaldarlo con evidencia, no con opinión. La "
            "diferencia entre un buen background y uno mediocre es la presencia de "
            "evidencia concreta."
        ),
        "elementos": [
            "Una estadística europea reciente (últimos 3 años) o referencia a un informe oficial (Comisión, Cedefop, Eurofound, agencia nacional).",
            "Una mención a documento de política EU relevante.",
            "Experiencia propia de las organizaciones del consorcio (datos reales).",
            "Identificación clara del gap: qué iniciativas existen pero son insuficientes y por qué.",
        ],
        "ejemplo_debil": "Existen muchas iniciativas pero no son suficientes para resolver este problema.",
        "ejemplo_fuerte": (
            "El Plan de Acción de Educación Digital de la Comisión Europea señala que menos "
            "del 30% de los educadores se sienten cómodos integrando herramientas digitales. "
            "Aunque existen programas nacionales de formación, nuestra experiencia con más "
            "de 100 docentes muestra que son mayoritariamente teóricos y no transferibles "
            "al aula."
        ),
        "evitar": [
            "Datos sin fuente o de hace más de 5 años",
            "'Hay muchos estudios que...' sin citar ninguno",
            "Confundir opinión con evidencia",
            "Listar referencias sin conectarlas al proyecto",
        ],
    },
    {
        "titulo": "Encaje con el propósito de la call",
        "prioridad": "alta",
        "obligatorio": "sí",
        "intencion": (
            "Demostrar explícitamente que el proyecto pertenece a ESTA call concreta — no "
            "a otra. La conexión nunca debe quedar implícita. Es la parte más débil de la "
            "mayoría de propuestas: tienen buena idea pero no la anclan a la call."
        ),
        "elementos": [
            "Identificar 2-3 objetivos centrales de la call.",
            "Para cada uno: qué significa + cómo el proyecto contribuye en la práctica.",
            "Usar lenguaje que refleje (sin copiar) el wording de la call.",
            "Reflejar el propósito específico de Small-scale Partnerships: ampliar acceso a newcomers y actores pequeños.",
        ],
        "ejemplo_debil": "Nuestro proyecto está alineado con todos los objetivos de la convocatoria.",
        "ejemplo_fuerte": (
            "La call enfatiza capacity building para organizaciones pequeñas y newcomers. "
            "Nuestro proyecto aborda esto formando a 4 clubes de menos de 50 socios en "
            "gestión de voluntariado y diseño de programas inclusivos — capacidades que "
            "ninguno tiene actualmente y que no podrían desarrollar individualmente."
        ),
        "evitar": [
            "Copy-paste de la Programme Guide",
            "Listar objetivos de la call sin conectar al proyecto",
            "'Estamos alineados con todas las prioridades'",
            "Más de 3 objetivos (pierde foco)",
        ],
    },
    {
        "titulo": "Objetivos generales orientados a cambio",
        "prioridad": "alta",
        "obligatorio": "sí",
        "intencion": (
            "Presentar entre 2 y 4 objetivos generales que describan un CAMBIO o MEJORA, "
            "no una actividad. Cada objetivo debe derivar visiblemente del problema "
            "definido en C1."
        ),
        "elementos": [
            "2-4 objetivos (más es contraproducente).",
            "Cada uno como statement orientado a resultado: 'Improve X among Y' / 'Strengthen capacity of Z'.",
            "Conexión explícita con el problema de C1.",
            "Lenguaje de outcome (improve, strengthen, enhance, build), no de actividad (organise, deliver, create).",
        ],
        "ejemplo_debil": "Organizar 4 workshops sobre coaching inclusivo para clubes deportivos.",
        "ejemplo_fuerte": (
            "Fortalecer la capacidad de 4 clubes pequeños para diseñar y mantener programas "
            "deportivos inclusivos para jóvenes en riesgo de abandono."
        ),
        "evitar": [
            "Actividades disfrazadas de objetivos",
            "Objetivos vagos ('mejorar el deporte en Europa')",
            "Más de 4 objetivos",
            "Objetivos no derivados del problema",
        ],
    },
    {
        "titulo": "EU added value y prioridades",
        "prioridad": "media",
        "obligatorio": "sí",
        "intencion": (
            "Conectar el proyecto a 1-2 prioridades EU (horizontales o sectoriales) y "
            "demostrar por qué requiere dimensión europea. Lo que la cooperación produce "
            "y no podría conseguir un solo país."
        ),
        "elementos": [
            "1-2 prioridades máximo (selectiva, no exhaustiva).",
            "Para cada una: por qué es relevante al problema y cómo el proyecto contribuye en la práctica.",
            "Límites de un enfoque puramente nacional.",
            "Beneficios concretos de la cooperación: intercambio de prácticas, aprendizaje comparativo, métodos transferibles.",
        ],
        "ejemplo_debil": (
            "El proyecto tiene dimensión europea porque los partners son de varios países "
            "y promueve la inclusión."
        ),
        "ejemplo_fuerte": (
            "El abandono deportivo adolescente afecta a países con sistemas formativos muy "
            "distintos: Italia aporta su modelo de scuola sportiva, Polonia su red de "
            "clubes municipales, España su experiencia con voluntarios. Sin esta diversidad, "
            "ningún partner podría desarrollar un protocolo de retención adaptable a "
            "contextos diversos."
        ),
        "evitar": [
            "'Promovemos la inclusión' sin demostrarlo",
            "Listar muchas prioridades sin justificación",
            "Mencionar partnership internacional sin explicar el valor",
            "Outcomes limitados al nivel local",
        ],
    },
]

for i, c in enumerate(criterios_ejemplo, 1):
    add_heading(doc, f"CRITERIO {i} — {c['titulo']}", level=3)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Prioridad: {c['prioridad']}    |    Obligatorio: {c['obligatorio']}")
    meta_run.italic = True
    meta_run.font.color.rgb = GREY
    meta_run.font.size = Pt(10)
    meta.paragraph_format.space_after = Pt(6)

    add_para(doc, "INTENCIÓN", bold=True, color=NAVY, size=10)
    add_para(doc, c["intencion"])

    add_para(doc, "ELEMENTOS", bold=True, color=NAVY, size=10)
    for e in c["elementos"]:
        add_bullet(doc, e)

    add_para(doc, "EJEMPLOS", bold=True, color=NAVY, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run("DÉBIL:  ")
    r1.bold = True
    r1.font.color.rgb = RED
    r1.font.size = Pt(10)
    r2 = p.add_run(c["ejemplo_debil"])
    r2.italic = True
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run("FUERTE: ")
    r1.bold = True
    r1.font.color.rgb = GREEN
    r1.font.size = Pt(10)
    r2 = p.add_run(c["ejemplo_fuerte"])
    r2.italic = True
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    add_para(doc, "EVITAR", bold=True, color=NAVY, size=10)
    for e in c["evitar"]:
        add_bullet(doc, e)

    if i < len(criterios_ejemplo):
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 60)
        sep_run.font.color.rgb = GREY
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Section 7: Checklist ──
add_heading(doc, "7. Checklist final antes de entregar", level=1)
add_para(doc, "Antes de mandar tu documento, revisa:", italic=True, color=GREY)

checklist = [
    "La pregunta tiene IDENTIFICACIÓN, CONTEXTO, CONEXIONES, REGLA GLOBAL (si aplica) y LÍMITE DE EXTENSIÓN.",
    "Hay entre 4 y 6 criterios (no menos, no más).",
    "Cada criterio tiene título corto, prioridad y obligatorio marcados.",
    "Cada criterio tiene los 4 bloques de contenido: INTENCIÓN, ELEMENTOS, EJEMPLOS, EVITAR.",
    "Cada criterio tiene un ejemplo DÉBIL y un ejemplo FUERTE concretos.",
    "Las prioridades suman sentido (no todos son 'alta', no todos son 'baja').",
    "Los ejemplos están en el idioma del formulario final.",
    "No hay copy-paste de la Programme Guide.",
    "Cada criterio = una idea distinta. Si dos se solapan, fusionar.",
]
for item in checklist:
    p = doc.add_paragraph()
    box = p.add_run("☐  ")
    box.font.size = Pt(12)
    box.font.color.rgb = NAVY
    text = p.add_run(item)
    text.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

# ── Section 8: Cómo entregar ──
add_heading(doc, "8. Cómo entregar el trabajo", level=1)
add_para(doc, "Un documento Word por pregunta, nombrado así:")
add_code_block(
    doc,
    "criterios_1.1.docx\n"
    "criterios_1.2.docx\n"
    "criterios_2.1.1.docx\n"
    "..."
)
add_para(
    doc,
    "Una vez entregado, el equipo técnico lo procesa y lo ingesta en el sistema. La IA del "
    "Writer empezará a usar estos criterios en su próxima generación.",
)

# ── Section 9: FAQ ──
add_heading(doc, "9. Dudas frecuentes", level=1)

faqs = [
    ("¿Y si no se me ocurre un ejemplo débil?",
     "Piensa en propuestas reales que has visto rechazadas o con baja puntuación. ¿Qué frase típica salía? Esa es tu ejemplo débil."),
    ("¿Puedo poner más de un ejemplo fuerte por criterio?",
     "Sí, si aportan ángulos distintos. Pero uno bueno > tres mediocres."),
    ("¿Y si dos criterios comparten un EVITAR?",
     "Si es realmente compartido por todos, súbelo a REGLA GLOBAL. Si solo dos, repítelo en ambos sin problema."),
    ("¿Cuánta extensión debe tener cada criterio?",
     "Entre 200 y 400 palabras en total (sumando los 4 bloques). Si pasa de 500, probablemente hay que dividirlo en dos criterios."),
    ("¿En qué idioma escribo los criterios?",
     "Los criterios (intención, elementos, etc.) en español. Los ejemplos débil/fuerte en el idioma del formulario final (suele ser inglés)."),
]
for q, a in faqs:
    add_para(doc, q, bold=True, color=NAVY, size=11)
    add_para(doc, a)
    doc.add_paragraph()

# Save
doc.save(OUTPUT_PATH)
print(f"OK: {OUTPUT_PATH}")
